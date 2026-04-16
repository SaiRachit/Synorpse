"""
File Creator - Create files and folders with AI-generated content
Supports creating Python, Word, PDF, Markdown, and text files with intelligent content
"""
import os
from pathlib import Path
from typing import Optional, Dict, Any, List
import logging
from datetime import datetime
from dotenv import dotenv_values
from groq import Groq

logger = logging.getLogger("file_creator")

env_vars = dotenv_values(".env")
GROQ_CONVERSATION_KEY = env_vars.get("GroqConversationManager")

groq_client = Groq(api_key=GROQ_CONVERSATION_KEY) if GROQ_CONVERSATION_KEY else None


class FileCreator:
    """Create files and folders with AI-generated content"""
    
    def __init__(self, chatbot_func, search_func):
        """
        Initialize FileCreator
        
        Args:
            chatbot_func: Async function to generate content via ChatBot (fallback)
            search_func: Async function to search web for information
        """
        self.chatbot = chatbot_func
        self.search = search_func
        self.groq_client = groq_client
        
        if not self.groq_client:
            logger.warning("GroqConversationManager API key not found, using fallback ChatBot")
    
    async def create_folder(self, path: str) -> Dict[str, Any]:
        """
        Create a folder
        
        Args:
            path: Path where to create folder
        
        Returns:
            Dict with success status and message
        """
        try:
            folder_path = Path(path)
            folder_path.mkdir(parents=True, exist_ok=True)
            
            return {
                "success": True,
                "message": f"Folder created: {folder_path}",
                "path": str(folder_path)
            }
        except Exception as e:
            logger.error(f"Failed to create folder: {e}", exc_info=True)
            return {
                "success": False,
                "message": f"Failed to create folder: {str(e)}",
                "path": path
            }
    
    async def generate_code_content(self, description: str, language: str = "python") -> str:
        """
        Generate code content using Groq API directly
        
        Args:
            description: What the code should do
            language: Programming language
        
        Returns:
            Generated code as string
        """
        prompt = f"""Write {language} code for: {description}

Requirements:
- Write clean, well-commented code
- Include docstrings/comments explaining the logic
- Follow best practices for {language}
- Make it production-ready
- ONLY return the code, no explanations or markdown formatting"""
        
        
        try:
            # Use Groq client directly for better quality
            if self.groq_client:
                response = self.groq_client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.3,
                    max_tokens=2000
                )
                code = response.choices[0].message.content
            else:
                # Fallback to ChatBot
                code = await self.chatbot(prompt)
        except Exception as e:
            logger.error(f"Groq API failed, using fallback: {e}")
            code = await self.chatbot(prompt)
        
        # Clean up markdown code blocks if present
        if "```" in code:
            # Extract code from markdown block
            lines = code.split("\n")
            in_code_block = False
            clean_lines = []
            for line in lines:
                if line.strip().startswith("```"):
                    in_code_block = not in_code_block
                    continue
                if in_code_block or not any(line.strip().startswith(x) for x in ["```", "#", "**"]):
                    clean_lines.append(line)
            code = "\n".join(clean_lines).strip()
        
        return code
    
    async def generate_document_content(self, topic: str, use_web: bool = True) -> str:
        """
        Generate document content about a topic using Groq API
        
        Args:
            topic: Topic to write about
            use_web: Whether to search web for information first
        
        Returns:
            Generated content as string
        """
        content_parts = []
        web_research = ""
        
        # Search web if requested
        if use_web:
            try:
                search_results = await self.search(f"detailed information about {topic}")
                web_research = f"\n\nWeb Research Results:\n{search_results}\n"
            except Exception as e:
                pass
        
        # Generate comprehensive content with Groq
        prompt = f"""Write a comprehensive, well-researched document about: {topic}
{web_research}
Requirements:
- Write in a professional, informative style
- Include relevant sections with clear headings (use ## for sections)
- Provide detailed, factual information
- Make it well-structured and easy to read
- Length: 800-1200 words
- Use markdown formatting for headings
- ONLY return the document content, no metadata or explanations"""
        
        try:
            # Use Groq client directly for better quality
            if self.groq_client:
                response = self.groq_client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.7,
                    max_tokens=3000
                )
                content = response.choices[0].message.content
            else:
                # Fallback to ChatBot
                content = await self.chatbot(prompt)
        except Exception as e:
            content = await self.chatbot(prompt)
        
        content_parts.append(content)
        
        return "\n".join(content_parts)
    
    def open_file_location(self, file_path: str) -> None:
        """
        Open File Explorer and select the created file
        
        Args:
            file_path: Path to the file to reveal
        """
        import subprocess
        try:
            # Open Explorer and select the file
            subprocess.run(['explorer', '/select,', str(file_path)], check=False)
        except Exception as e:
            pass
    
    
    async def create_python_file(self, path: str, description: str) -> Dict[str, Any]:
        """
        Create a Python file with generated code
        
        Args:
            path: Path where to create file
            description: What the code should do
        
        Returns:
            Dict with success status and message
        """
        try:
            # Generate code
            code = await self.generate_code_content(description, "python")
            
            # Ensure .py extension
            file_path = Path(path)
            if file_path.suffix != ".py":
                file_path = file_path.with_suffix(".py")
            
            # Create parent directories
            file_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Write file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(code)
            
            return {
                "success": True,
                "message": f"Python file created: {file_path}",
                "path": str(file_path),
                "content_preview": code[:200] + "..." if len(code) > 200 else code
            }
        except Exception as e:
            logger.error(f"Failed to create Python file: {e}", exc_info=True)
            return {
                "success": False,
                "message": f"Failed to create Python file: {str(e)}",
                "path": path
            }
    
    async def create_text_file(self, path: str, topic: str, use_web: bool = True) -> Dict[str, Any]:
        """
        Create a text/markdown file with generated content
        
        Args:
            path: Path where to create file
            topic: Topic to write about
            use_web: Whether to search web for information
        
        Returns:
            Dict with success status and message
        """
        try:
            # Generate content
            content = await self.generate_document_content(topic, use_web)
            
            # Add metadata header
            header = f"""# {topic}
Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

---

"""
            
            full_content = header + content
            
            # Ensure correct extension
            file_path = Path(path)
            if file_path.suffix not in [".txt", ".md"]:
                file_path = file_path.with_suffix(".txt")
            
            # Create parent directories
            file_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Write file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(full_content)
            
            return {
                "success": True,
                "message": f"Text file created: {file_path}",
                "path": str(file_path),
                "content_preview": full_content[:200] + "..." if len(full_content) > 200 else full_content
            }
        except Exception as e:
            logger.error(f"Failed to create text file: {e}", exc_info=True)
            return {
                "success": False,
                "message": f"Failed to create text file: {str(e)}",
                "path": path
            }
    
    async def create_word_file(self, path: str, topic: str, use_web: bool = True) -> Dict[str, Any]:
        """
        Create a Word document with generated content and relevant images
        
        Args:
            path: Path where to create file
            topic: Topic to write about
            use_web: Whether to search web for information
        
        Returns:
            Dict with success status and message
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError:
            return {
                "success": False,
                "message": "python-docx not installed.",
                "path": path
            }
        
        try:
            # Search for relevant images
            images_to_embed = []
            try:
                from InternetImages import search_google_images, HEADERS
                import requests
                import io
                from PIL import Image as PILImage
                
                logger.info(f"Searching for images related to: {topic}")
                image_urls = search_google_images(topic, num=3)
                
                # Download and validate images
                for url in image_urls[:2]:  # Limit to 2 images for documents
                    try:
                        response = requests.get(url, headers=HEADERS, timeout=10)
                        if response.status_code == 200 and len(response.content) > 100:
                            img = PILImage.open(io.BytesIO(response.content))
                            # Convert to RGB if needed
                            if img.mode in ('RGBA', 'LA', 'P'):
                                rgb_img = PILImage.new('RGB', img.size, (255, 255, 255))
                                if img.mode == 'P':
                                    img = img.convert('RGBA')
                                if img.mode in ('RGBA', 'LA'):
                                    rgb_img.paste(img, mask=img.split()[-1])
                                else:
                                    rgb_img.paste(img)
                                img = rgb_img
                            elif img.mode != 'RGB':
                                img = img.convert('RGB')
                            
                            images_to_embed.append(img)
                    except Exception as e:
                        logger.debug(f"Could not load image from {url}: {e}")
                        continue
                
                logger.info(f"Loaded {len(images_to_embed)} images for document")
            except Exception as e:
                logger.warning(f"Image search failed: {e}")
            
            # Generate content
            content = await self.generate_document_content(topic, use_web)
            
            # Create document
            doc = Document()
            
            # Add title
            title = doc.add_heading(topic, 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add metadata
            meta = doc.add_paragraph()
            meta.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
            meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            doc.add_paragraph()  # Spacer
            
            # Add first image if available
            if images_to_embed:
                try:
                    # Save first image temporarily
                    temp_img_path = "temp_doc_img.jpg"
                    images_to_embed[0].save(temp_img_path, 'JPEG')
                    doc.add_picture(temp_img_path, width=Inches(5))
                    doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    doc.add_paragraph()  # Spacer
                    import os
                    os.remove(temp_img_path)
                except Exception as e:
                    logger.warning(f"Could not embed first image: {e}")
            
            # Parse and add content
            lines = content.split('\n')
            section_count = 0
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check if it's a heading
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                    section_count += 1
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                    section_count += 1
                    
                    # Add second image after first major section if available
                    if section_count == 2 and len(images_to_embed) > 1:
                        try:
                            temp_img_path = "temp_doc_img2.jpg"
                            images_to_embed[1].save(temp_img_path, 'JPEG')
                            doc.add_picture(temp_img_path, width=Inches(4.5))
                            doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            doc.add_paragraph()
                            import os
                            os.remove(temp_img_path)
                        except Exception as e:
                            logger.warning(f"Could not embed second image: {e}")
                            
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                else:
                    # Regular paragraph
                    doc.add_paragraph(line)
            
            # Ensure .docx extension
            file_path = Path(path)
            if file_path.suffix != ".docx":
                file_path = file_path.with_suffix(".docx")
            
            # Create parent directories
            file_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Save document
            doc.save(file_path)
            
            # Save document
            doc.save(file_path)
            
            # Open file location in Explorer
            self.open_file_location(str(file_path))
            
            img_msg = f" with {len(images_to_embed)} images" if images_to_embed else ""
            return {
                "success": True,
                "message": f"Word document created: {file_path}{img_msg}",
                "path": str(file_path),
                "content_preview": content[:200] + "..." if len(content) > 200 else content
            }
        except Exception as e:
            logger.error(f"Failed to create Word file: {e}", exc_info=True)
            return {
                "success": False,
                "message": f" Failed to create Word file: {str(e)}",
                "path": path
            }
    
    async def create_pdf_file(self, path: str, topic: str, use_web: bool = True) -> Dict[str, Any]:
        """
        Create a PDF document with generated content and relevant images
        
        Args:
            path: Path where to create file
            topic: Topic to write about
            use_web: Whether to search web for information
        
        Returns:
            Dict with success status and message
        """
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image as RLImage
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
        except ImportError:
            return {
                "success": False,
                "message": "reportlab not installed.",
                "path": path
            }
        
        try:
            # Search for relevant images
            images_to_embed = []
            try:
                from InternetImages import search_google_images, HEADERS
                import requests
                import io
                from PIL import Image as PILImage
                
                logger.info(f"Searching for images related to: {topic}")
                image_urls = search_google_images(topic, num=3)
                
                # Download and validate images
                for url in image_urls[:2]:  # Limit to 2 images for PDFs
                    try:
                        response = requests.get(url, headers=HEADERS, timeout=10)
                        if response.status_code == 200 and len(response.content) > 100:
                            img = PILImage.open(io.BytesIO(response.content))
                            # Convert to RGB if needed
                            if img.mode in ('RGBA', 'LA', 'P'):
                                rgb_img = PILImage.new('RGB', img.size, (255, 255, 255))
                                if img.mode == 'P':
                                    img = img.convert('RGBA')
                                if img.mode in ('RGBA', 'LA'):
                                    rgb_img.paste(img, mask=img.split()[-1])
                                else:
                                    rgb_img.paste(img)
                                img = rgb_img
                            elif img.mode != 'RGB':
                                img = img.convert('RGB')
                            
                            # Save temporarily for PDF embedding
                            temp_path = f"temp_pdf_img_{len(images_to_embed)}.jpg"
                            img.save(temp_path, 'JPEG')
                            images_to_embed.append(temp_path)
                    except Exception as e:
                        logger.debug(f"Could not load image from {url}: {e}")
                        continue
                
                logger.info(f"Loaded {len(images_to_embed)} images for PDF")
            except Exception as e:
                logger.warning(f"Image search failed: {e}")
            
            # Generate content
            content = await self.generate_document_content(topic, use_web)
            
            # Ensure .pdf extension
            file_path = Path(path)
            if file_path.suffix != ".pdf":
                file_path = file_path.with_suffix(".pdf")
            
            # Create parent directories
            file_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Create PDF
            doc = SimpleDocTemplate(str(file_path), pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            # Title style
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor='darkblue',
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            # Add title
            story.append(Paragraph(topic, title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Add metadata
            meta_style = ParagraphStyle(
                'Meta',
                parent=styles['Normal'],
                fontSize=10,
                textColor='gray',
                alignment=TA_CENTER
            )
            story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", meta_style))
            story.append(Spacer(1, 0.3*inch))
            
            # Add first image if available
            if images_to_embed:
                try:
                    img = RLImage(images_to_embed[0], width=5*inch, height=3*inch)
                    story.append(img)
                    story.append(Spacer(1, 0.3*inch))
                except Exception as e:
                    logger.warning(f"Could not embed first image: {e}")
            
            # Add content
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['BodyText'],
                fontSize=11,
                alignment=TA_JUSTIFY,
                spaceAfter=12
            )
            
            # Parse content
            paragraphs = content.split('\n\n')
            para_count = 0
            for para in paragraphs:
                para = para.strip()
                if para:
                    # Escape XML special characters
                    para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(para, body_style))
                    story.append(Spacer(1, 0.1*inch))
                    para_count += 1
                    
                    # Add second image after a few paragraphs
                    if para_count == 3 and len(images_to_embed) > 1:
                        try:
                            img = RLImage(images_to_embed[1], width=4.5*inch, height=2.7*inch)
                            story.append(img)
                            story.append(Spacer(1, 0.3*inch))
                        except Exception as e:
                            logger.warning(f"Could not embed second image: {e}")
            
            # Build PDF
            doc.build(story)
            
            # Clean up temporary image files
            for temp_img in images_to_embed:
                try:
                    import os
                    os.remove(temp_img)
                except:
                    pass
            
            # Build PDF
            doc.build(story)
            
            # Open file location in Explorer
            self.open_file_location(str(file_path))
            
            img_msg = f" with {len(images_to_embed)} images" if images_to_embed else ""
            return {
                "success": True,
                "message": f"PDF document created: {file_path}{img_msg}",
                "path": str(file_path),
                "content_preview": content[:200] + "..." if len(content) > 200 else content
            }
        except Exception as e:
            logger.error(f"Failed to create PDF file: {e}", exc_info=True)
            return {
                "success": False,
                "message": f" Failed to create PDF file: {str(e)}",
                "path": path
            }
    
    async def create_from_conversation(self, path: str, file_type: str, conversation_context: List[Dict], topic: Optional[str] = None) -> Dict[str, Any]:
        """
        Create a file from conversation context
        
        Args:
            path: Path where to create file
            file_type: Type of file (word, pdf, text, markdown)
            conversation_context: Recent conversation turns
            topic: Optional topic override
        
        Returns:
            Dict with success status and message
        """
        try:
            # Extract conversation content
            conversation_text = "\n\n".join([
                f"User: {turn.get('user_input', '')}\nAssistant: {turn.get('assistant_response', '')}"
                for turn in conversation_context[-5:]  # Last 5 turns
            ])
            
            # Determine topic from conversation if not provided
            if not topic:
                topic = "Summary of our conversation"
                if conversation_context:
                    # Try to extract main topic
                    first_user = conversation_context[0].get('user_input', '')
                    if first_user:
                        topic = f"Conversation about {first_user[:50]}"
            
            # Generate summary/document from conversation
            prompt = f"""Based on the following conversation, create a comprehensive document.

Conversation:
{conversation_text}

Create a well-structured document that:
- Summarizes the key points discussed
- Organizes information in logical sections
- Adds relevant context and explanations
- Makes it professional and informative

Title: {topic}

Return only the document content with markdown formatting."""
            
            logger.info(f"Creating {file_type} file from conversation context")
            
            # Get content from AI
            if self.groq_client:
                response = self.groq_client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.7,
                    max_tokens=3000
                )
                content = response.choices[0].message.content
            else:
                content = await self.chatbot(prompt)
            
            # Create appropriate file type
            if file_type.lower() in ["word", "doc", "docx"]:
                return await self._create_word_from_content(path, topic, content)
            elif file_type.lower() == "pdf":
                return await self._create_pdf_from_content(path, topic, content)
            else:
                return await self._create_text_from_content(path, topic, content)
                
        except Exception as e:
            return {
                "success": False,
                "message": f"Failed to create file from conversation: {str(e)}",
                "path": path
            }
    
    async def _create_text_from_content(self, path: str, topic: str, content: str) -> Dict[str, Any]:
        """Helper to create text file from content"""
        file_path = Path(path)
        if file_path.suffix not in [".txt", ".md"]:
            file_path = file_path.with_suffix(".txt")
        
        header = f"""# {topic}
Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

---

"""
        file_path.parent.mkdir(parents=True, exist_ok=True)
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(header + content)
        
        self.open_file_location(str(file_path))
        return {
            "success": True,
            "message": f"Text file created: {file_path}",
            "path": str(file_path)
        }
    
    async def _create_word_from_content(self, path: str, topic: str, content: str) -> Dict[str, Any]:
        """Helper to create Word file from content"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            
            doc = Document()
            title = doc.add_heading(topic, 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            meta = doc.add_paragraph()
            meta.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
            meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph()
            
            for line in content.split('\n'):
                line = line.strip()
                if line:
                    if line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('# '):
                        doc.add_heading(line[2:], level=1)
                    else:
                        doc.add_paragraph(line)
            
            file_path = Path(path).with_suffix(".docx")
            file_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(file_path)
            
            self.open_file_location(str(file_path))
            return {
                "success": True,
                "message": f"Word document created: {file_path}",
                "path": str(file_path)
            }
        except Exception as e:
            return {
                "success": False,
                "message": f"Failed to create Word file: {str(e)}",
                "path": path
            }
    
    async def _create_pdf_from_content(self, path: str, topic: str, content: str) -> Dict[str, Any]:
        """Helper to create PDF file from content"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
            
            file_path = Path(path).with_suffix(".pdf")
            file_path.parent.mkdir(parents=True, exist_ok=True)
            
            doc = SimpleDocTemplate(str(file_path), pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=24, alignment=TA_CENTER)
            story.append(Paragraph(topic, title_style))
            story.append(Spacer(1, 0.3*inch))
            
            body_style = ParagraphStyle('CustomBody', parent=styles['BodyText'], fontSize=11, alignment=TA_JUSTIFY)
            for para in content.split('\n\n'):
                if para.strip():
                    para_clean = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(para_clean, body_style))
                    story.append(Spacer(1, 0.1*inch))
            
            doc.build(story)
            self.open_file_location(str(file_path))
            return {
                "success": True,
                "message": f"PDF document created: {file_path}",
                "path": str(file_path)
            }
        except Exception as e:
            return {
                "success": False,
                "message": f"Failed to create PDF file: {str(e)}",
                "path": path
            }




# Global instance
_file_creator = None

def get_file_creator(chatbot_func=None, search_func=None):
    """Get global file creator instance"""
    global _file_creator
    if _file_creator is None and chatbot_func and search_func:
        _file_creator = FileCreator(chatbot_func, search_func)
    return _file_creator
